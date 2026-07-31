"""Build the Intro Connect QA contractor package.

Produces a single self-contained zip:

    QA-INTRO-CONNECT-PACKAGE.zip
      README-package.txt
      QA-TESTER-HANDOUT.docx
      QA-TEST-SCRIPT.xlsx
      sample-media/            every file the upload rows ask for

Replaces the old Agenda Builder only package (scripts/qa-agenda-package.py).
Phase A is carried over intact as regression coverage; phases B through H cover
the event loop that shipped after it was written: agenda conversion, the .ics
download, the host guest list import, announcements, surveys, the cross event
directory, the attendee email scoping, and the upgrade page.

The tester needs no credentials from us. Every account in the script is a plus
alias of one Gmail address the tester controls, so invitation emails, welcome
emails and temporary passwords all land in their own inbox, and the email leg
of each flow gets tested in the same pass. Nothing in this package grants
access to anything; there is nothing here to rotate.

Run from the repo root with the backend venv (needs openpyxl, python-docx,
Pillow; tooling dependencies only, deliberately NOT in requirements.txt):

    backend/.venv/Scripts/python.exe scripts/qa-package.py
"""
from __future__ import annotations

import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation
from PIL import Image, ImageDraw

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "build" / "qa-package"
MEDIA = OUT / "sample-media"

APP_URL = "https://app.intro-connect.com"
TOOL_URL = f"{APP_URL}/agenda"
BUILDER_URL = f"{APP_URL}/agenda/new"

HEADER = [
    "Row",
    "Area",
    "What to do",
    "What should happen",
    "Pass/Fail",
    "What actually happened",
    "Screenshot file(s)",
]

# ---------------------------------------------------------------------------
# PHASE A: the Agenda Builder, public, no account. Carried over from the
# original package as regression coverage; the tool has not changed shape.
# ---------------------------------------------------------------------------
ROWS_A = [
    # --- Landing page -------------------------------------------------------
    ("A1", "Landing", f"Open {TOOL_URL} in Chrome on a desktop computer.",
     "Page loads. You see a badge reading 'Free tool, no account needed', a heading "
     "'Build an event agenda in minutes', and a blue 'Start building an agenda' button."),
    ("A2", "Landing", "Read the three cards below the heading.",
     "Three cards: 'Add your sessions', 'Download a Word agenda', 'Turn it into an event'. "
     "No spelling mistakes, nothing cut off."),
    ("A3", "Landing", "In the top nav, click 'Pricing'. Then come back and click 'Blog', then 'News'.",
     "Each one takes you to intro-connect.com (NOT app.intro-connect.com). Pricing lands on the "
     "pricing section. Blog and News show real articles, not an error page."),
    ("A4", "Landing", "In the top nav, click 'Log in'. Go back, then click 'Start for free'.",
     "'Log in' opens the Intro Connect login page. 'Start for free' opens the sign up page. "
     "Neither one 404s and neither leaves the app and bounces back."),
    ("A5", "Landing", "Scroll to the very bottom and read the footer.",
     "Footer has Features, Pricing, FAQ, Blog, News, Privacy, Terms, Contact, and the line "
     "'HOST BETTER. CONNECT DEEPER. BUILD WHAT MATTERS.' It must NOT say 'Front Range Dev Co'."),
    ("A6", "Landing", "Click 'Start building an agenda'.",
     "You land on the builder at /agenda/new."),

    # --- Event details ------------------------------------------------------
    ("A7", "Details", "Look at the top of the builder page.",
     "Heading 'Agenda Builder' with a blue 'Free' badge next to it, and text on the right "
     "reading 'Saves as you type'."),
    ("A8", "Details", "Type 'Denver Founders Dinner' into Event name.",
     "Text appears. Within a second or two the top right changes from 'Saves as you type' to "
     "'Saved' with a time."),
    ("A9", "Details", "Set Start date to a date about a month from today. Leave End date empty.",
     "Date is accepted. The hint under End date says it can be left empty for a single day event."),
    ("A10", "Details", "Set Start time 17:30 and End time 21:00. Type 'The Loft' into Venue name.",
     "All three accepted, nothing resets or clears."),
    ("A11", "Details", "Click 'Add description, organizer, and logo'.",
     "The section expands to reveal Description, Address, Virtual event link, Organizer name, "
     "Organizer company, Organizer email, Event website, and Event logo. Clicking it again collapses it."),
    ("A12", "Details", "Fill in a Description, Organizer name, Organizer company, Organizer email, "
     "and Event website (use intro-connect.com).",
     "All fields accept text and stay filled."),

    # --- Logo upload --------------------------------------------------------
    ("A13", "Logo", "Under Event logo click 'Upload a logo' and pick sample-media\\logo-square.png.",
     "A small square preview of the logo appears with a 'Remove' link next to it."),
    ("A14", "Logo", "Click 'Remove' next to the logo preview.",
     "Preview disappears and the 'Upload a logo' button comes back."),
    ("A15", "Logo", "Upload sample-media\\logo-wide.jpg.",
     "Accepted, preview appears. JPEG is a supported format."),
    ("A16", "Logo", "Remove it, then upload sample-media\\logo-transparent.webp.",
     "Accepted, preview appears. WebP is a supported format."),
    ("A17", "Logo", "Remove it, then try to upload sample-media\\oversized-logo.png (this one is over 1 MB).",
     "REFUSED with a readable message saying the file is too large and logos can be up to 1 MB. "
     "No preview appears. The page must not freeze or go blank."),
    ("A18", "Logo", "Try to upload sample-media\\not-an-image.png. (It has a .png name but is really a text file.)",
     "REFUSED with a readable message asking for a PNG, JPEG, or WebP image. No preview. "
     "Nothing crashes."),
    ("A19", "Logo", "Upload sample-media\\logo-square.png again and leave it in place for the rest of the test.",
     "Preview appears and stays."),

    # --- Sessions -----------------------------------------------------------
    ("A20", "Sessions", "Look at the Agenda section before adding anything.",
     "You see a dashed box saying 'Add your first session to get started' with a blue "
     "'Add your first session' button. It should be obvious what to do next."),
    ("A21", "Sessions", "Click 'Add your first session'. Fill it in: date = your event date, "
     "Start 17:30, End 18:15, Title 'Doors open and welcome drinks', Room 'Rooftop terrace', "
     "Speaker blank, Description 'Name tags, a drink, and no programming.'",
     "A session card appears with all those fields and keeps everything you typed."),
    ("A22", "Sessions", "Click 'Add session' and add a second: same date, Start 18:15, End 18:45, "
     "Title 'Opening remarks', Speaker 'Scott Weiss', Room 'Main room'.",
     "Second card appears below the first."),
    ("A23", "Sessions", "Add a third: same date, Start 18:45, End 20:00, Title 'Dinner'.",
     "Three session cards, all under one date heading showing the full date "
     "(for example 'Saturday, August 1, 2026')."),
    ("A24", "Sessions", "On the second session, click the copy icon (Duplicate).",
     "An identical copy appears directly below it. You now have four sessions. Editing the copy "
     "must NOT change the original."),
    ("A25", "Sessions", "Delete the duplicate using the bin icon.",
     "It disappears, the other three are untouched. Back to three sessions."),
    ("A26", "Sessions", "In the Title box of session one, type a very long title (paste the same "
     "sentence about 10 times).",
     "The box accepts it and the card does not break the page layout or push anything off screen."),

    # --- Reordering ---------------------------------------------------------
    ("A27", "Reorder", "On the LAST session, click the up chevron once.",
     "That session visibly swaps places with the one above it. The order on screen actually changes."),
    ("A28", "Reorder", "On the FIRST session, look at the up chevron. On the LAST, look at the down chevron.",
     "The first session's up arrow is greyed out and does nothing. The last session's down arrow "
     "is greyed out and does nothing."),
    ("A29", "Reorder", "Grab the dotted grip handle on the LEFT of the first session card and drag it "
     "down below the last one, then let go.",
     "The card follows your mouse while dragging, and drops into the new position. The order sticks."),
    ("A30", "Reorder", "Click into a session's Title box and try to select the text by dragging across it.",
     "You can select and edit text normally. Dragging inside a text box must NOT pick up and move "
     "the whole card."),
    ("A31", "Reorder", "Deliberately put the sessions in an order that does NOT match their times "
     "(for example put the 20:00 dinner first). Then click Preview, then back to Edit.",
     "YOUR order is kept exactly as you arranged it. The tool must NOT quietly re-sort them by time."),
    ("A32", "Reorder", "Now click 'Sort by time' next to the date heading.",
     "Only now do the sessions rearrange into time order, earliest first."),

    # --- Validation and warnings -------------------------------------------
    ("A33", "Validation", "On any session set the End time EARLIER than the Start time "
     "(for example Start 18:00, End 17:00).",
     "The End box turns red and a red message appears saying the end time needs to come after "
     "the start time."),
    ("A34", "Validation", "With that bad time still in place, click 'Download Word Agenda'.",
     "It REFUSES and shows a message saying some sessions end before they start. No file downloads."),
    ("A35", "Validation", "Fix the time so it is valid again.",
     "The red warning disappears."),
    ("A36", "Validation", "Make two sessions overlap on purpose (both 18:00 to 19:00).",
     "An AMBER (orange) warning appears saying they overlap and that is fine if the tracks run at "
     "once. This is a warning only, it must NOT block anything."),
    ("A37", "Validation", "With the overlap still there, click 'Download Word Agenda'.",
     "The download DOES work. Overlaps never block. Then undo the overlap."),
    ("A38", "Validation", "Add a new session but leave its Date empty. Click 'Download Word Agenda'.",
     "It refuses with a message saying every session needs a date. The undated session shows under "
     "a heading 'Date not set'. Give it a date afterwards."),

    # --- Multi day ----------------------------------------------------------
    ("A39", "Multi day", "Set the End date to the day after your start date. Add two more sessions "
     "and set their Date to that second day.",
     "Sessions group under TWO separate date headings, earlier day first, each showing the full date."),

    # --- Preview and autosave ----------------------------------------------
    ("A40", "Preview", "Click the 'Preview' tab.",
     "A clean document style preview: logo at the top, event name, the date range, the venue, your "
     "description, then each day with its sessions and times, then an Organizer block, then a small "
     "line 'Agenda created with Intro Connect | intro-connect.com'."),
    ("A41", "Preview", "Compare the preview against what you typed in Edit.",
     "Everything matches: same sessions, same order, same times, same speakers and rooms."),
    ("A42", "Autosave", "Press F5 to reload the page. Do not save anything first.",
     "Everything is still there: event details, logo, and every session in the same order. "
     "Nothing is lost."),
    ("A43", "Autosave", "Close the browser tab completely, open a new tab, and go back to "
     f"{BUILDER_URL}",
     "Your agenda is still there."),

    # --- The Word document --------------------------------------------------
    ("A44", "Download", "Go back to the Edit tab and click 'Download Word Agenda'.",
     "A .docx file downloads. The filename is based on your event name, "
     "for example 'denver-founders-dinner-agenda.docx'."),
    ("A45", "Download", "Open the downloaded file in Microsoft Word.",
     "It opens as a normal Word document. NOT a 'this file is corrupt' warning, NOT a web page."),
    ("A46", "Download", "Check the top of the document.",
     "Your logo (centred), the event name as a large heading, the date line, and the venue line."),
    ("A47", "Download", "Check the body of the document.",
     "Your description, then each day as its own heading, then a table with the time on the left "
     "and the session title, speaker, room and description on the right."),
    ("A48", "Download", "Check the session order in the document.",
     "It matches the order you arranged on screen, exactly. Not re-sorted."),
    ("A49", "Download", "Check the bottom of the document and the page footer.",
     "An 'Organizer' block with your name, company, email and website. In the page footer: "
     "'Agenda created with Intro Connect | intro-connect.com'. There must be NO adverts or "
     "large promotional blocks anywhere in the document."),
    ("A50", "Download", "Now type some new text into the document in Word, and change a session title.",
     "You can edit it freely and save it. It is a real editable Word file."),
    ("A51", "Privacy", "Go back to the builder. Put the text 'SECRET-NOTE-123' into the "
     "'Private notes' box on any session. Download the agenda again and open it in Word. "
     "Use Word's Find (Ctrl+F) to search for SECRET-NOTE-123.",
     "It is NOT found anywhere in the document. Private notes must never appear in the download. "
     "This one matters, please be thorough."),
    ("A52", "Download", "In the builder, add a session Link (use https://example.com/session). "
     "Download again and check that session in the document.",
     "The link text appears under the session."),

    # --- Conversion CTA -----------------------------------------------------
    ("A53", "CTA", "Look at the page straight after a successful download.",
     "A panel appears: 'Your agenda is ready.' with text about turning it into an event, a blue "
     "'Create Your Event' button and a 'Download Word Agenda' button."),
    ("A54", "CTA", "Click the secondary 'Download Word Agenda' button in that panel.",
     "The file downloads again."),
    ("A55", "CTA", "Click 'Create Your Event' (you are not logged in yet).",
     "You are taken to the sign up page. IMPORTANT: do NOT sign up yet. That is row B1, and it "
     "continues from exactly this point. For now press Back and confirm your agenda is still there."),

    # --- Empty and edge cases ----------------------------------------------
    ("A56", "Edge", "Open a private/incognito window and go to "
     f"{BUILDER_URL} . Without typing anything, click 'Download Word Agenda'.",
     "A document downloads. Opened in Word it says 'Untitled event' and "
     "'No sessions have been added yet.' It must not crash or produce a broken file."),
    ("A57", "Edge", "Still in the incognito window, paste this into a session Link box: "
     "javascript:alert(1) . Then add a date and title and download the file.",
     "The document downloads normally and that text does NOT appear as a clickable link in it. "
     "Nothing pops up. Report exactly what you see in the document. Then close the incognito window."),

    # --- Responsive ---------------------------------------------------------
    ("A58", "Mobile", f"On a real phone, open {TOOL_URL}",
     "Landing page fits the screen. No sideways scrolling. Text is readable without zooming."),
    ("A59", "Mobile", "On the phone, tap the hamburger (three lines) menu in the top right.",
     "A menu opens with How it works, Features, Pricing, FAQ, Agenda Builder, Blog, News, Log in, "
     "Start for free. Tapping one closes the menu and goes there."),
    ("A60", "Mobile", "On the phone, build a small agenda: event name, one date, two sessions.",
     "All boxes are tappable, the keyboard does not cover what you are typing, and nothing "
     "overlaps or runs off the side of the screen."),
    ("A61", "Mobile", "On the phone, use the up and down chevrons to reorder the two sessions.",
     "They reorder correctly by tapping. (Dragging on a touch screen may be fiddly, that is what "
     "the arrows are for. Note in your answer whether dragging worked on your phone.)"),
    ("A62", "Mobile", "On the phone, download the agenda.",
     "The .docx downloads to the phone and can be opened."),
    ("A63", "Browsers", "Repeat rows A21, A29 and A44 in a second browser "
     "(Firefox or Safari, whichever you did not use).",
     "Adding a session, dragging to reorder, and downloading all behave the same as in Chrome."),
    ("A64", "Desktop", "On a monitor at least 1024 pixels wide, look at the top nav of "
     "intro-connect.com with the window at full width, then slowly narrow the window.",
     "At full width every nav link sits on ONE row with nothing wrapped onto a second line and no "
     "hamburger. As the window narrows past roughly 1024 pixels the links collapse into the "
     "hamburger. At no width does anything overlap or fall off the edge."),
]

# ---------------------------------------------------------------------------
# PHASE B: sign up through the agenda, and the event it creates. HOST ONE.
# ---------------------------------------------------------------------------
ROWS_B = [
    ("B1", "Sign up", "Back in your main browser window (where your Phase A agenda lives), click "
     "'Create Your Event' again. This time sign up: email = your Gmail with +ichost1 added (see "
     "the handout), any password you like, and a made up name like 'Jordan Wells'. Complete the "
     "profile step with made up details.",
     "Sign up works. After the profile step you land on a CONFIRMATION screen that shows the "
     "agenda you built in Phase A: event name, dates, venue. Nothing was lost during sign up."),
    ("B2", "Convert", "On that confirmation screen, click Create.",
     "You land on the page of your new event. The name, date and venue match your agenda, and "
     "the AGENDA ITSELF is displayed on the event page: every session, in your order, with times."),
    ("B3", "Email", "Check the +ichost1 inbox (it is your own Gmail inbox).",
     "A welcome email from Intro Connect arrived. Note how long it took. If nothing arrives "
     "within 10 minutes, that is a FAIL for this row; carry on regardless."),
    ("B4", "Calendar", "On the event page, click 'Add to calendar'.",
     "An .ics file downloads. Opening it in your calendar app shows the event with the right "
     "title on the right day. A one day event appears on that single day, not spread across two."),
    ("B5", "Limits", "Go to your events list and try to create a SECOND event.",
     "REFUSED with a readable message that the free plan includes 1 event and to upgrade to host "
     "more. No second event appears."),
]

# ---------------------------------------------------------------------------
# PHASE C: the host guest list import. Still HOST ONE, on Event One.
# ---------------------------------------------------------------------------
ROWS_C = [
    ("C1", "Import", "On the event page click 'Import guests'. In the modal, PASTE the two lines "
     "below (replace YOURADDRESS with your real Gmail name first):\n\n"
     "email,name,role,company\n"
     "YOURADDRESS+icguest1@gmail.com,Riley Park,Founder,Parkline",
     "The modal reads it and shows '1 guest ready to import' with no skipped rows. The Import "
     "button shows 'Import 1'."),
    ("C2", "Import", "Click Import.",
     "A result screen: '1 added to this event', '1 new account created, 0 already had one', and "
     "a line explaining everyone new gets an invitation email with their own sign in link. "
     "AT NO POINT is a password shown to you. If you can see a password anywhere on this screen, "
     "screenshot it and mark this row FAIL."),
    ("C3", "Import", "Open sample-media\\guest-list-TEMPLATE.csv in Notepad. Replace every "
     "YOURADDRESS with your real Gmail name and save. Back in the app, click 'Import guests' "
     "again and this time use 'Upload CSV' to upload that file.",
     "The preview shows '2 guests ready to import' and '2 rows will be skipped', naming them by "
     "line: Line 4 (not a valid email address) and Line 5 (phone is not 10 digits). The two good "
     "rows are NOT blocked by the two bad ones."),
    ("C4", "Import", "Click Import, then look at the attendee list on the event page.",
     "Both new guests appear. THE IMPORTANT ONE: the guest whose name in the file was "
     "\"Doe, Jane\" (with a comma, in quotes) shows as exactly 'Doe, Jane' at company "
     "'Acme, Inc.' The comma survived and nothing shifted into the wrong box."),
    ("C5", "Import", "Import the same pasted line from C1 again.",
     "Nothing duplicates. The result says 0 new accounts created and 1 already had one. The "
     "attendee list does not grow."),
    ("C6", "Email", "Check your inbox for the three +ic guest aliases.",
     "Each one received an invitation email naming the event, containing a TEMPORARY PASSWORD "
     "and an 'Open your directory' button. Keep these emails, the next phase uses them."),
    ("C7", "Capacity", "Look at the attendee count line on the event page (host view).",
     "It shows the current count against a limit of 50 (the free plan cap), for example "
     "'4 of 50'. The number matches how many guests are actually listed, plus nobody twice."),
]

# ---------------------------------------------------------------------------
# PHASE D: the guest side. GUEST ONE in a SECOND browser profile.
# ---------------------------------------------------------------------------
ROWS_D = [
    ("D1", "Guest", "Open a SECOND browser profile or a different browser entirely (not just a "
     "new tab; the two accounts must not share a login). Open the +icguest1 invitation email, "
     f"click 'Open your directory', and sign in with that alias and the temporary password "
     "from the email.",
     "You can sign in. If you are asked to finish a profile first, fill it with made up details. "
     "You end up able to see Event One."),
    ("D2", "Guest", "Look over the event page as this guest.",
     "You can see: the event details, the agenda with every session, and the attendee grid. You "
     "can NOT see: any 'Import guests' button, any 'Post an announcement' button, or any survey "
     "editing controls. The page is read and participate, not manage."),
    ("D3", "Privacy", "In the attendee grid, open another guest's card (for example Doe, Jane).",
     "The profile opens with name, role, company. There is NO email address anywhere on it. "
     "A guest must never see another guest's email address. If you can find one anywhere as this "
     "account, screenshot it and mark this row FAIL. This is one of the rows that matters most."),
    ("D4", "Privacy", "Still signed in as the guest, click your own name or profile area.",
     "Your own details are visible to you and editable. Your own email showing to YOU is fine."),
]

# ---------------------------------------------------------------------------
# PHASE E: announcements. HOST ONE and GUEST ONE, both browsers side by side.
# ---------------------------------------------------------------------------
ROWS_E = [
    ("E1", "Announce", "As HOST ONE on the event page, click 'Post an announcement'. Title "
     "'Room change', body 'Dinner has moved to the rooftop. Lift to floor 6.' Click Post.",
     "The announcement appears immediately at the TOP of the event page, above the agenda, with "
     "your host name and the time."),
    ("E2", "Announce", "As GUEST ONE, reload the event page.",
     "The announcement is there with a small NEW badge. Reload the page a second time: the badge "
     "is gone (it was new to you once). The announcement itself stays."),
    ("E3", "Announce", "As HOST ONE, under the announcement click 'Also send by email'.",
     "Your computer's mail app opens a draft: the guest addresses are in BCC (NOT in To), and "
     "the subject is the announcement title. CLOSE THE DRAFT WITHOUT SENDING. If the button "
     "instead says 'Copy N addresses for BCC', click it and confirm the addresses land on your "
     "clipboard; that is the intended behaviour for long guest lists."),
    ("E4", "Announce", "As HOST ONE, try to post an announcement with an EMPTY body.",
     "The Post button is disabled while the body is empty. You cannot post nothing."),
    ("E5", "Announce", "As HOST ONE, delete the 'Room change' announcement (bin icon).",
     "A styled confirmation dialog appears (from the app, not a bare browser popup) warning it "
     "cannot be undone. Confirm. It disappears. As GUEST ONE, reload: it is gone there too."),
]

# ---------------------------------------------------------------------------
# PHASE F: the survey. HOST ONE writes it, GUEST ONE answers it.
# ---------------------------------------------------------------------------
ROWS_F = [
    ("F1", "Survey", "As HOST ONE on the event page, find the Survey section and click "
     "'Add a survey'. Fill in questions 1 and 2 ('How was the venue?', 'How useful were the "
     "introductions?') but leave question 3 EMPTY. Click Save survey.",
     "Refused with a readable message that every question needs some text. Nothing is saved."),
    ("F2", "Survey", "Fill in question 3 ('Would you come again?') and Save.",
     "Saved. As the host you now see a RESULTS view: '0 responses' and 'no answers' beside each "
     "question. You do NOT get an answering form; the host reads results, guests answer."),
    ("F3", "Survey", "As GUEST ONE, reload the event page and find the Survey section.",
     "The three questions appear, each with five round buttons 1 to 5 and a note '1 low, 5 high'. "
     "Below them: 'Your host sees the totals for everyone, not who said what.'"),
    ("F4", "Survey", "As GUEST ONE, try to Submit with only two questions answered.",
     "Refused with a message to answer all three questions."),
    ("F5", "Survey", "Answer all three (pick 5, 4, 3) and Submit.",
     "'Thanks for the feedback'. The form stays with your answers selected and the button now "
     "reads 'Update answers'."),
    ("F6", "Survey", "Change your first answer from 5 to 4 and click 'Update answers'.",
     "Accepted. This must count as EDITING your response, not voting twice."),
    ("F7", "Survey", "As HOST ONE, reload the event page and read the results.",
     "'1 response'. The averages match what the guest picked last (4, 4, 3), each question shows "
     "a small bar chart of 1 to 5, and NOWHERE does it say who answered. If any name or email "
     "appears next to an answer, screenshot it and mark this row FAIL."),
    ("F8", "Survey", "As HOST ONE, click 'Edit questions'. Read the warning, then change question "
     "1 to 'How was the food?' and Save.",
     "Before saving, a note warns that changing the wording clears the answers given so far. "
     "After saving, results show '0 responses' again. As GUEST ONE, reload: the survey asks "
     "fresh with the new wording and your old answers are gone. (Answers belong to the question "
     "as it was asked; keeping them would report replies to a question nobody was asked.)"),
]

# ---------------------------------------------------------------------------
# PHASE G: the cross event directory. Needs HOST TWO and GUEST TWO, so that
# two people exist who have NEVER been in the same room.
# ---------------------------------------------------------------------------
ROWS_G = [
    ("G1", "Setup", "In a THIRD browser profile (or log out of Host One first), go to "
     f"{APP_URL} and sign up a new account with your +ichost2 alias and a made up name. Create "
     "an event the ordinary way (no agenda this time): name 'Boulder Product Night', a date "
     "next month.",
     "Account and event both created. This is Event Two, and its host has no connection to "
     "Event One."),
    ("G2", "Setup", "As HOST TWO, use 'Import guests' to paste one row (replace YOURADDRESS):\n\n"
     "email,name,role,company\n"
     "YOURADDRESS+icguest2@gmail.com,Sam Alder,Engineer,Alderworks",
     "Imports cleanly, invitation email arrives at the +icguest2 alias."),
    ("G3", "Setup", "In a FOURTH browser profile, sign in as GUEST TWO using that invitation "
     "email's temporary password.",
     "You are in, and can see Event Two. Guest Two and Guest One have now never shared an event."),
    ("G4", "Directory", "As GUEST TWO on the Event Two page, find the 'Cross event directory' "
     "card near the bottom.",
     "It is there, the switch is OFF by default, and the text says people from other events "
     "could find your profile and message you, and that your email address is never shown. "
     "Nobody is listed anywhere without turning this on themselves."),
    ("G5", "Directory", "As GUEST TWO, open 'Directory' in the top nav.",
     "The directory page loads with a banner explaining YOU are not listed yet, so you cannot "
     "message people here and they cannot find you. Guest One does NOT appear (they have not "
     "opted in either). The page is essentially empty."),
    ("G6", "Directory", "Switch to GUEST ONE's browser. On the Event One page, turn the "
     "'Cross event directory' switch ON.",
     "The switch flips and a message confirms you are now listed."),
    ("G7", "Directory", "Back as GUEST TWO, reload the Directory page.",
     "Guest One (Riley Park) now appears as a card: name, role, company. NO email address "
     "anywhere on the card. The banner still says you yourself are not listed."),
    ("G8", "Directory", "As GUEST TWO, open Riley Park's card and try to send them a message.",
     "It does NOT work: the app refuses, because messaging through the directory needs BOTH "
     "people listed, and you are not. An error or a 'not found' style message is acceptable "
     "here; silently pretending to send is not. Report exactly what you see."),
    ("G9", "Directory", "As GUEST TWO, go to Event Two and turn the 'Cross event directory' "
     "switch ON. Then return to the Directory and message Riley Park: 'Hello from another "
     "event entirely.'",
     "The message SENDS now that both of you are listed."),
    ("G10", "Directory", "As GUEST ONE, open Messages.",
     "The message from Sam Alder is there. Reply to it. As GUEST TWO, confirm the reply "
     "arrives. Two people who never attended anything together are now talking, because both "
     "chose to be findable."),
    ("G11", "Directory", "As GUEST ONE, turn the directory switch OFF again on Event One. As "
     "GUEST TWO, reload the Directory page.",
     "Riley Park no longer appears in the directory."),
    ("G12", "Directory", "As HOST ONE (who created Event One but never joined it as a guest), "
     "find the 'Cross event directory' card on the event page and try to switch it on.",
     "Refused with a readable message that you are not on this event's guest list, so there is "
     "nothing to list. The switch returns to off. Hosting an event does not put you in the "
     "directory; only being a guest somewhere can."),
]

# ---------------------------------------------------------------------------
# PHASE H: what the host can see, the upgrade page, and the phone.
# ---------------------------------------------------------------------------
ROWS_H = [
    ("H1", "Host view", "As HOST ONE, open a guest's card from the Event One attendee grid.",
     "As the HOST you DO see the guest's email address on their profile. The host owns their "
     "guest list; fellow guests (row D3) do not. Both halves have to hold."),
    ("H2", "Billing", "As HOST ONE, open the upgrade page (the app will have offered it at row "
     "B5; it is also at /upgrade).",
     "Three plans: Free, Starter, Pro. A monthly/annual toggle. Monthly shows $39 and $99; "
     "annual shows $390 and $990 marked as billed yearly. The numbers must be exactly these."),
    ("H3", "Billing", "Click 'Choose Starter' with ANNUAL selected.",
     "A Stripe payment page opens showing Intro Connect and $390.00 per year. STOP THERE: "
     "close the tab. DO NOT enter card details, DO NOT pay. The row passes if the Stripe page "
     "opened with the right amount."),
    ("H4", "Mobile", "On a real phone, sign in as GUEST ONE and open the Event One page.",
     "Everything from this test renders on the phone: announcements area, agenda, survey with "
     "its 1 to 5 buttons, the directory switch, and the attendee grid. No sideways scrolling, "
     "nothing overlapping, everything tappable."),
    ("H5", "Tidy up", "In every browser profile you used, sign out.",
     "Done. Leave the accounts and events in place; do not delete anything. Scott cleans up "
     "server side after reviewing your results."),
]

PHASES = [
    ("PHASE A  —  the Agenda Builder, no account needed.  (64 rows)", ROWS_A),
    ("PHASE B  —  sign up and turn the agenda into an event.  HOST ONE.  (5 rows)", ROWS_B),
    ("PHASE C  —  import a guest list.  HOST ONE.  (7 rows)", ROWS_C),
    ("PHASE D  —  the guest side.  GUEST ONE, second browser.  (4 rows)", ROWS_D),
    ("PHASE E  —  announcements.  Both browsers.  (5 rows)", ROWS_E),
    ("PHASE F  —  the survey.  Both browsers.  (8 rows)", ROWS_F),
    ("PHASE G  —  the cross event directory.  All four accounts.  (12 rows)", ROWS_G),
    ("PHASE H  —  host view, billing page, phone.  (5 rows)", ROWS_H),
]


def build_media() -> None:
    MEDIA.mkdir(parents=True, exist_ok=True)

    square = Image.new("RGB", (400, 400), (37, 99, 235))
    d = ImageDraw.Draw(square)
    d.ellipse((90, 90, 310, 310), fill=(255, 255, 255))
    square.save(MEDIA / "logo-square.png")

    wide = Image.new("RGB", (600, 200), (13, 27, 42))
    d = ImageDraw.Draw(wide)
    d.rectangle((40, 60, 560, 140), fill=(79, 141, 247))
    wide.save(MEDIA / "logo-wide.jpg", format="JPEG", quality=90)

    trans = Image.new("RGBA", (300, 300), (0, 0, 0, 0))
    d = ImageDraw.Draw(trans)
    d.ellipse((30, 30, 270, 270), fill=(37, 99, 235, 255))
    trans.save(MEDIA / "logo-transparent.webp", format="WEBP")

    # Comfortably over the 1 MB cap: random noise does not compress away.
    import os

    noise = Image.frombytes("RGB", (900, 900), os.urandom(900 * 900 * 3))
    noise.save(MEDIA / "oversized-logo.png", format="PNG", compress_level=0)

    (MEDIA / "not-an-image.png").write_text(
        "This is a plain text file that has been given a .png name on purpose.\n"
        "It is here so you can check the site refuses files that are not really images.\n",
        encoding="utf-8",
    )

    # The import template for phase C. Line numbers are load bearing: row C3
    # tells the tester exactly which lines get skipped, so anything added here
    # must keep the bad email on line 4 and the bad phone on line 5.
    (MEDIA / "guest-list-TEMPLATE.csv").write_text(
        "email,name,role,company,industry,phone\n"
        "YOURADDRESS+icg3@gmail.com,Avery Stone,Operations,Northwind,Software,(303) 555-0161\n"
        'YOURADDRESS+icg4@gmail.com,"Doe, Jane",Partner,"Acme, Inc.",Consulting,303-555-0162\n'
        "notanemail,Broken Row,Tester,Nowhere,Testing,\n"
        "YOURADDRESS+icg5@gmail.com,Sam Field,Designer,Atelier,Design,555-123\n",
        encoding="utf-8",
    )

    size_mb = (MEDIA / "oversized-logo.png").stat().st_size / (1024 * 1024)
    assert size_mb > 1, f"oversized-logo.png is only {size_mb:.2f} MB, it must exceed 1 MB"


def build_workbook() -> None:
    wb = Workbook()

    ws = wb.active
    ws.title = "READ ME"
    lines = [
        ("Intro Connect: full product QA", True),
        ("", False),
        ("Fill in ONLY the three yellow columns on the Test Script sheet:", False),
        ("  Pass/Fail (dropdown: PASS / FAIL / BLOCKED)", False),
        ("  What actually happened (one or two sentences; exact error text if any)", False),
        ("  Screenshot file(s) (filename in your screenshots folder, e.g. G7.png)", False),
        ("Everything else is read only reference. Work top to bottom, keep the order.", False),
        ("", False),
        ("PASS means it did exactly what the 'What should happen' column says.", False),
        ("FAIL means it did something else. Say what it actually did.", False),
        ("BLOCKED means you could not attempt the row. Say why.", False),
        ("", False),
        ("Example of a filled row:", True),
        ('  A29 | FAIL | "Card followed the mouse but snapped back to its old place on release." | A29.png',
         False),
        ("", False),
        ("Screenshot every FAIL. Screenshots are optional on a PASS unless the row asks.", False),
        ("", False),
        ("Before starting phase B, read the Accounts section of the handout. The whole", False),
        ("script uses plus aliases of ONE Gmail address you control, and phases D to G", False),
        ("need separate browser profiles so the accounts stay signed in side by side.", False),
        ("", False),
        ("Full instructions and the rules: QA-TESTER-HANDOUT.docx in this package.", False),
    ]
    for i, (text, bold) in enumerate(lines, start=1):
        cell = ws.cell(row=i, column=1, value=text)
        cell.font = Font(bold=bold, size=12 if bold and i == 1 else 11)
    ws.column_dimensions["A"].width = 105

    sh = wb.create_sheet("Test Script")
    head_fill = PatternFill("solid", fgColor="0D1B2A")
    fill_fill = PatternFill("solid", fgColor="FFF3C4")
    phase_fill = PatternFill("solid", fgColor="E8EEF9")

    for ci, name in enumerate(HEADER, start=1):
        c = sh.cell(row=1, column=ci, value=name)
        c.font = Font(bold=True, color="FFFFFF")
        c.fill = head_fill
        c.alignment = Alignment(vertical="center", wrap_text=True)

    r = 2

    def add_banner(text: str):
        nonlocal r
        c = sh.cell(row=r, column=1, value=text)
        c.font = Font(bold=True)
        c.fill = phase_fill
        sh.merge_cells(start_row=r, start_column=1, end_row=r, end_column=len(HEADER))
        r += 1

    def add_rows(rows):
        nonlocal r
        for row_id, area, todo, expect in rows:
            for ci, value in enumerate([row_id, area, todo, expect], start=1):
                c = sh.cell(row=r, column=ci, value=value)
                c.alignment = Alignment(vertical="top", wrap_text=True)
            for ci in range(5, len(HEADER) + 1):
                c = sh.cell(row=r, column=ci)
                c.fill = fill_fill
                c.alignment = Alignment(vertical="top", wrap_text=True)
            r += 1

    for banner, rows in PHASES:
        add_banner(banner)
        add_rows(rows)

    widths = [7, 13, 62, 62, 11, 42, 20]
    for ci, w in enumerate(widths, start=1):
        sh.column_dimensions[get_column_letter(ci)].width = w
    sh.freeze_panes = "A2"

    dv = DataValidation(type="list", formula1='"PASS,FAIL,BLOCKED"', allow_blank=True)
    sh.add_data_validation(dv)
    dv.add(f"E2:E{r - 1}")

    sh.auto_filter.ref = f"A1:{get_column_letter(len(HEADER))}{r - 1}"

    wb.save(OUT / "QA-TEST-SCRIPT.xlsx")


def _p(doc, text="", *, bold=False, size=11, color=None, space_after=6, align=None):
    para = doc.add_paragraph()
    if align:
        para.alignment = align
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return para


def build_handout() -> None:
    doc = Document()
    doc.core_properties.title = "Intro Connect QA"

    _p(doc, "Intro Connect", bold=True, size=20, color=RGBColor(0x1D, 0x4E, 0xD8), space_after=0)
    _p(doc, "Full product QA: manual test", bold=True, size=15, space_after=14)

    _p(doc, "What this is", bold=True, size=13)
    _p(doc, "Intro Connect helps event organizers turn an event into a private directory of "
            "everyone who attended, so people can actually find each other afterwards.")
    _p(doc, "You are testing the whole loop: the free public Agenda Builder, signing up, turning "
            "an agenda into a real event, importing a guest list, host announcements, a guest "
            "survey, and a directory that connects people across different events. Most of this "
            "shipped very recently and has never been used by anyone except the people who built "
            "it. That is your job.", space_after=14)

    _p(doc, "Where to go", bold=True, size=13)
    _p(doc, TOOL_URL, bold=True)
    _p(doc, "Phase A starts there. The rest of the product is at app.intro-connect.com and the "
            "script walks you into it at the right moment.", space_after=14)

    _p(doc, "Accounts: read this before phase B", bold=True, size=13)
    _p(doc, "You need ONE Gmail address you control. Gmail treats yourname+anything@gmail.com as "
            "your own inbox, and the script uses that: every account you create is a plus alias "
            "of your address, so every email the product sends lands where you can check it. "
            "That is deliberate, because checking those emails is part of the test.")
    _p(doc, "You will end up with four accounts, created at the moment the script asks:")
    accounts = [
        "HOST ONE  =  yourname+ichost1@gmail.com  (you sign up; row B1)",
        "GUEST ONE  =  yourname+icguest1@gmail.com  (created by the import; row C1)",
        "HOST TWO  =  yourname+ichost2@gmail.com  (you sign up; row G1)",
        "GUEST TWO  =  yourname+icguest2@gmail.com  (created by the import; row G2)",
    ]
    for a in accounts:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(4)
        para.add_run(a).font.size = Pt(11)
    _p(doc, "Phases D to G need two accounts signed in AT THE SAME TIME, side by side. Use "
            "separate browser profiles (Chrome: the avatar menu, 'Add profile') or different "
            "browsers. Plain new tabs share a login and will not work.")
    _p(doc, "Use made up names and details on every profile. The plus aliases are yours; "
            "everything else about these people is fiction.", space_after=14)

    _p(doc, "The job", bold=True, size=13)
    _p(doc, "Open QA-TEST-SCRIPT.xlsx and work through the Test Script tab top to bottom, in "
            "order. Fill in the three yellow columns as you go. Rows build on each other; "
            "skipping one usually breaks the ones after it.")
    _p(doc, "Estimated time: 8 to 12 hours. Fine to split across days; your work stays where "
            "you left it.", bold=True, space_after=14)

    _p(doc, "The rules that matter", bold=True, size=13)
    rules = [
        "This is the live production site. The accounts and events you create are expected and "
        "wanted, but keep everything inside the script: made up people, made up events, and only "
        "your own plus aliased email address anywhere an email goes.",
        "NEVER enter card details and never pay for anything. One row (H3) opens a Stripe "
        "payment page on purpose; it ends with you closing the tab.",
        "One row (E3) opens an email draft in your mail app; it ends with you closing the draft "
        "without sending.",
        "All uploads come from the sample-media folder. Do not upload your own photos or any "
        "real company's logo.",
        "Do not delete accounts or events when you finish. Scott reviews first, then cleans up.",
        "If something is unclear, or you hit an error you cannot get past, or you see anything "
        "that looks like a REAL person's data (a name or email that is not one of yours), stop "
        "and message Scott rather than improvising.",
        "Everything in this package, and everything you see on the site, is confidential. Do not "
        "share it, post it, or feed it into anything.",
    ]
    for rule in rules:
        para = doc.add_paragraph(style="List Number")
        para.paragraph_format.space_after = Pt(6)
        para.add_run(rule).font.size = Pt(11)
    doc.add_paragraph()

    _p(doc, "Screenshots", bold=True, size=13)
    _p(doc, "Make a folder called screenshots next to the workbook. Name each file after its row: "
            "G7.png, and G7-2.png if you need a second. Screenshot every FAIL. On a PASS a "
            "screenshot is optional unless the row asks.", space_after=14)

    _p(doc, "What good feedback looks like", bold=True, size=13)
    _p(doc, "Not helpful:  \"Doesn't work\"", space_after=2)
    _p(doc, "Helpful:  \"Clicked Import on the pasted row. The spinner ran about five seconds, "
            "then the modal just closed. No result screen, no error, and the guest did not "
            "appear in the list. Same thing on a second try. Screenshot G-import-2.png is the "
            "network tab showing a 500.\"", space_after=14)

    _p(doc, "A few things worth knowing", bold=True, size=13)
    notes = [
        "Guests never pay and guests never see adverts. If anything ever asks a GUEST account "
        "for money, that is a serious FAIL.",
        "A guest must never be able to see another guest's email address (rows D3 and G7). The "
        "HOST of an event does see their own guest list's addresses (row H1). Both are intended; "
        "the difference is the point.",
        "Nobody appears in the cross event directory unless they turned it on themselves, and "
        "messaging through it only works when BOTH people are listed. Several rows check "
        "each side of this.",
        "Sessions in an agenda stay in the order YOU put them; there is a separate Sort by time "
        "button. Private notes on a session must never appear in the downloaded document (A51).",
        "The survey is deliberately exactly three questions answered 1 to 5. Changing a "
        "question's wording clears its answers on purpose (row F8 explains why).",
        "Emails can take a few minutes. Note arrival times; if one never arrives, that is a "
        "finding, record it.",
    ]
    for note in notes:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(6)
        para.add_run(note).font.size = Pt(11)
    doc.add_paragraph()

    _p(doc, "When you are done", bold=True, size=13)
    _p(doc, "Zip the filled in workbook together with your screenshots folder and send it back to "
            "Scott the same way you received this package.")

    doc.save(OUT / "QA-TESTER-HANDOUT.docx")


README = f"""INTRO CONNECT: QA PACKAGE — START HERE

1. Open QA-TESTER-HANDOUT.docx and read it first. It explains what you are
   testing, the four plus aliased accounts you will create, and the rules
   that matter (no card details, nothing real except your own email).

2. Open QA-TEST-SCRIPT.xlsx. Start at the READ ME tab, then work through the
   Test Script tab top to bottom, in order. Phases build on each other.

3. sample-media\\ holds every file the rows ask you to upload, plus
   guest-list-TEMPLATE.csv which you edit once (your own Gmail name in place
   of YOURADDRESS) before the import rows use it.

4. Create a folder called "screenshots" next to the workbook. Save every
   screenshot there, named after the row (G7.png, G7-2.png, ...).

5. When you are done: zip the filled workbook + your screenshots folder and
   send it back to Scott the same way you received this package.

Phase A starts here, and nothing on the public site links to it yet:
   {TOOL_URL}

You need ONE Gmail address you control. Everything else is created as you go,
exactly when the script says to.

Questions or anything unexpected (an error you cannot get past, anything that
looks like a real person's data): stop and contact Scott before continuing.

This package is confidential. Do not share any part of it.
"""


def main() -> None:
    if OUT.exists():
        shutil.rmtree(OUT)
    OUT.mkdir(parents=True)

    build_media()
    build_workbook()
    build_handout()
    (OUT / "README-package.txt").write_text(README, encoding="utf-8")

    zip_path = ROOT / "build" / "QA-INTRO-CONNECT-PACKAGE.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
        for path in sorted(OUT.rglob("*")):
            if path.is_file():
                z.write(path, path.relative_to(OUT))

    total = sum(len(rows) for _, rows in PHASES)
    print(f"wrote {zip_path} ({zip_path.stat().st_size / 1024:.0f} KB)")
    print(f"  test rows: {total} across {len(PHASES)} phases "
          f"({' + '.join(str(len(r)) for _, r in PHASES)})")
    for path in sorted(OUT.rglob("*")):
        if path.is_file():
            print(f"  {path.relative_to(OUT)}  ({path.stat().st_size / 1024:.1f} KB)")


if __name__ == "__main__":
    main()
