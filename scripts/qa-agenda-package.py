"""Build the Agenda Builder QA contractor package.

Produces a single self-contained zip:

    QA-AGENDA-BUILDER-PACKAGE.zip
      README-package.txt
      QA-TESTER-HANDOUT.docx
      QA-TEST-SCRIPT.xlsx
      sample-media/            every file the upload rows ask for

Mirrors the structure of the Rally UP contractor package so a tester who has
worked one recognises the other.

Run from the repo root with the backend venv (needs openpyxl, python-docx,
Pillow). openpyxl is a tooling dependency only and is deliberately NOT in
backend/requirements.txt:

    backend/.venv/Scripts/python.exe scripts/qa-agenda-package.py
"""
from __future__ import annotations

import io
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation
from PIL import Image, ImageDraw

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "build" / "qa-agenda"
MEDIA = OUT / "sample-media"

TOOL_URL = "https://app.intro-connect.com/agenda"
BUILDER_URL = "https://app.intro-connect.com/agenda/new"

HEADER = [
    "Row",
    "Area",
    "What to do",
    "What should happen",
    "Pass/Fail",
    "What actually happened",
    "Screenshot file(s)",
]
FILL_COLS = {"Pass/Fail", "What actually happened", "Screenshot file(s)"}

# ---------------------------------------------------------------------------
# The test script. Phase A needs no account at all: the Agenda Builder is a
# public tool, which is why the tester is never given a production login.
# ---------------------------------------------------------------------------
ROWS = [
    # --- Landing page -------------------------------------------------------
    ("A1", "Landing", f"Open {TOOL_URL} in Chrome on a desktop computer.",
     "Page loads. You see a badge reading 'Free tool, no account needed', a heading "
     "'Build an event agenda in minutes', and a blue 'Start building an agenda' button."),
    ("A2", "Landing", "Read the three cards below the heading.",
     "Three cards: 'Add your sessions', 'Download a Word agenda', 'Turn it into an event'. "
     "No spelling mistakes, nothing cut off."),
    ("A3", "Landing", "In the top nav, click 'Pricing'. Then come back and click 'Blog', then 'News'.",
     "Each one takes you to intro-connect.com (NOT app.intro-connect.com). Pricing lands on the "
     "pricing section. Blog and News show real articles, not an error page."),
    ("A4", "Landing", "In the top nav, click 'Log in'. Go back, then click 'Start for free'.",
     "'Log in' opens the Intro Connect login page. 'Start for free' opens the sign up page. "
     "Neither one 404s and neither leaves the app and bounces back."),
    ("A5", "Landing", "Scroll to the very bottom and read the footer.",
     "Footer has Features, Pricing, FAQ, Blog, News, Privacy, Terms, Contact, and the line "
     "'HOST BETTER. CONNECT DEEPER. BUILD WHAT MATTERS.' It must NOT say 'Front Range Dev Co'."),
    ("A6", "Landing", "Click 'Start building an agenda'.",
     "You land on the builder at /agenda/new."),

    # --- Event details ------------------------------------------------------
    ("A7", "Details", "Look at the top of the builder page.",
     "Heading 'Agenda Builder' with a blue 'Free' badge next to it, and text on the right "
     "reading 'Saves as you type'."),
    ("A8", "Details", "Type 'Denver Founders Dinner' into Event name.",
     "Text appears. Within a second or two the top right changes from 'Saves as you type' to "
     "'Saved' with a time."),
    ("A9", "Details", "Set Start date to a date about a month from today. Leave End date empty.",
     "Date is accepted. The hint under End date says it can be left empty for a single day event."),
    ("A10", "Details", "Set Start time 17:30 and End time 21:00. Type 'The Loft' into Venue name.",
     "All three accepted, nothing resets or clears."),
    ("A11", "Details", "Click 'Add description, organizer, and logo'.",
     "The section expands to reveal Description, Address, Virtual event link, Organizer name, "
     "Organizer company, Organizer email, Event website, and Event logo. Clicking it again collapses it."),
    ("A12", "Details", "Fill in a Description, Organizer name, Organizer company, Organizer email, "
     "and Event website (use intro-connect.com).",
     "All fields accept text and stay filled."),

    # --- Logo upload --------------------------------------------------------
    ("A13", "Logo", "Under Event logo click 'Upload a logo' and pick sample-media\\logo-square.png.",
     "A small square preview of the logo appears with a 'Remove' link next to it."),
    ("A14", "Logo", "Click 'Remove' next to the logo preview.",
     "Preview disappears and the 'Upload a logo' button comes back."),
    ("A15", "Logo", "Upload sample-media\\logo-wide.jpg.",
     "Accepted, preview appears. JPEG is a supported format."),
    ("A16", "Logo", "Remove it, then upload sample-media\\logo-transparent.webp.",
     "Accepted, preview appears. WebP is a supported format."),
    ("A17", "Logo", "Remove it, then try to upload sample-media\\oversized-logo.png (this one is over 1 MB).",
     "REFUSED with a readable message saying the file is too large and logos can be up to 1 MB. "
     "No preview appears. The page must not freeze or go blank."),
    ("A18", "Logo", "Try to upload sample-media\\not-an-image.png. (It has a .png name but is really a text file.)",
     "REFUSED with a readable message asking for a PNG, JPEG, or WebP image. No preview. "
     "Nothing crashes."),
    ("A19", "Logo", "Upload sample-media\\logo-square.png again and leave it in place for the rest of the test.",
     "Preview appears and stays."),

    # --- Sessions -----------------------------------------------------------
    ("A20", "Sessions", "Look at the Agenda section before adding anything.",
     "You see a dashed box saying 'Add your first session to get started' with a blue "
     "'Add your first session' button. It should be obvious what to do next."),
    ("A21", "Sessions", "Click 'Add your first session'. Fill it in: date = your event date, "
     "Start 17:30, End 18:15, Title 'Doors open and welcome drinks', Room 'Rooftop terrace', "
     "Speaker blank, Description 'Name tags, a drink, and no programming.'",
     "A session card appears with all those fields and keeps everything you typed."),
    ("A22", "Sessions", "Click 'Add session' and add a second: same date, Start 18:15, End 18:45, "
     "Title 'Opening remarks', Speaker 'Scott Weiss', Room 'Main room'.",
     "Second card appears below the first."),
    ("A23", "Sessions", "Add a third: same date, Start 18:45, End 20:00, Title 'Dinner'.",
     "Three session cards, all under one date heading showing the full date "
     "(for example 'Saturday, August 1, 2026')."),
    ("A24", "Sessions", "On the second session, click the copy icon (Duplicate).",
     "An identical copy appears directly below it. You now have four sessions. Editing the copy "
     "must NOT change the original."),
    ("A25", "Sessions", "Delete the duplicate using the bin icon.",
     "It disappears, the other three are untouched. Back to three sessions."),
    ("A26", "Sessions", "In the Title box of session one, type a very long title (paste the same "
     "sentence about 10 times).",
     "The box accepts it and the card does not break the page layout or push anything off screen."),

    # --- Reordering ---------------------------------------------------------
    ("A27", "Reorder", "On the LAST session, click the up chevron once.",
     "That session visibly swaps places with the one above it. The order on screen actually changes."),
    ("A28", "Reorder", "On the FIRST session, look at the up chevron. On the LAST, look at the down chevron.",
     "The first session's up arrow is greyed out and does nothing. The last session's down arrow "
     "is greyed out and does nothing."),
    ("A29", "Reorder", "Grab the dotted grip handle on the LEFT of the first session card and drag it "
     "down below the last one, then let go.",
     "The card follows your mouse while dragging, and drops into the new position. The order sticks."),
    ("A30", "Reorder", "Click into a session's Title box and try to select the text by dragging across it.",
     "You can select and edit text normally. Dragging inside a text box must NOT pick up and move "
     "the whole card."),
    ("A31", "Reorder", "Deliberately put the sessions in an order that does NOT match their times "
     "(for example put the 20:00 dinner first). Then click Preview, then back to Edit.",
     "YOUR order is kept exactly as you arranged it. The tool must NOT quietly re-sort them by time."),
    ("A32", "Reorder", "Now click 'Sort by time' next to the date heading.",
     "Only now do the sessions rearrange into time order, earliest first."),

    # --- Validation and warnings -------------------------------------------
    ("A33", "Validation", "On any session set the End time EARLIER than the Start time "
     "(for example Start 18:00, End 17:00).",
     "The End box turns red and a red message appears saying the end time needs to come after "
     "the start time."),
    ("A34", "Validation", "With that bad time still in place, click 'Download Word Agenda'.",
     "It REFUSES and shows a message saying some sessions end before they start. No file downloads."),
    ("A35", "Validation", "Fix the time so it is valid again.",
     "The red warning disappears."),
    ("A36", "Validation", "Make two sessions overlap on purpose (both 18:00 to 19:00).",
     "An AMBER (orange) warning appears saying they overlap and that is fine if the tracks run at "
     "once. This is a warning only, it must NOT block anything."),
    ("A37", "Validation", "With the overlap still there, click 'Download Word Agenda'.",
     "The download DOES work. Overlaps never block. Then undo the overlap."),
    ("A38", "Validation", "Add a new session but leave its Date empty. Click 'Download Word Agenda'.",
     "It refuses with a message saying every session needs a date. The undated session shows under "
     "a heading 'Date not set'. Give it a date afterwards."),

    # --- Multi day ----------------------------------------------------------
    ("A39", "Multi day", "Set the End date to the day after your start date. Add two more sessions "
     "and set their Date to that second day.",
     "Sessions group under TWO separate date headings, earlier day first, each showing the full date."),

    # --- Preview and autosave ----------------------------------------------
    ("A40", "Preview", "Click the 'Preview' tab.",
     "A clean document style preview: logo at the top, event name, the date range, the venue, your "
     "description, then each day with its sessions and times, then an Organizer block, then a small "
     "line 'Agenda created with Intro Connect | intro-connect.com'."),
    ("A41", "Preview", "Compare the preview against what you typed in Edit.",
     "Everything matches: same sessions, same order, same times, same speakers and rooms."),
    ("A42", "Autosave", "Press F5 to reload the page. Do not save anything first.",
     "Everything is still there: event details, logo, and every session in the same order. "
     "Nothing is lost."),
    ("A43", "Autosave", "Close the browser tab completely, open a new tab, and go back to "
     f"{BUILDER_URL}",
     "Your agenda is still there."),

    # --- The Word document --------------------------------------------------
    ("A44", "Download", "Go back to the Edit tab and click 'Download Word Agenda'.",
     "A .docx file downloads. The filename is based on your event name, "
     "for example 'denver-founders-dinner-agenda.docx'."),
    ("A45", "Download", "Open the downloaded file in Microsoft Word.",
     "It opens as a normal Word document. NOT a 'this file is corrupt' warning, NOT a web page."),
    ("A46", "Download", "Check the top of the document.",
     "Your logo (centred), the event name as a large heading, the date line, and the venue line."),
    ("A47", "Download", "Check the body of the document.",
     "Your description, then each day as its own heading, then a table with the time on the left "
     "and the session title, speaker, room and description on the right."),
    ("A48", "Download", "Check the session order in the document.",
     "It matches the order you arranged on screen, exactly. Not re-sorted."),
    ("A49", "Download", "Check the bottom of the document and the page footer.",
     "An 'Organizer' block with your name, company, email and website. In the page footer: "
     "'Agenda created with Intro Connect | intro-connect.com'. There must be NO adverts or "
     "large promotional blocks anywhere in the document."),
    ("A50", "Download", "Now type some new text into the document in Word, and change a session title.",
     "You can edit it freely and save it. It is a real editable Word file."),
    ("A51", "Privacy", "Go back to the builder. Put the text 'SECRET-NOTE-123' into the "
     "'Private notes' box on any session. Download the agenda again and open it in Word. "
     "Use Word's Find (Ctrl+F) to search for SECRET-NOTE-123.",
     "It is NOT found anywhere in the document. Private notes must never appear in the download. "
     "This one matters, please be thorough."),
    ("A52", "Download", "In the builder, add a session Link (use https://example.com/session). "
     "Download again and check that session in the document.",
     "The link text appears under the session."),

    # --- Conversion CTA -----------------------------------------------------
    ("A53", "CTA", "Look at the page straight after a successful download.",
     "A panel appears: 'Your agenda is ready.' with text about turning it into an event, a blue "
     "'Create Your Event' button and a 'Download Word Agenda' button."),
    ("A54", "CTA", "Click the secondary 'Download Word Agenda' button in that panel.",
     "The file downloads again."),
    ("A55", "CTA", "Click 'Create Your Event' (you are not logged in).",
     "You are taken to the sign up page. IMPORTANT: do not create an account, just confirm where "
     "it goes, then press Back."),
    ("A56", "CTA", "After pressing Back, check your agenda.",
     "Your agenda is still there and nothing has been lost."),

    # --- Empty and edge cases ----------------------------------------------
    ("A57", "Edge", "Open a private/incognito window and go to "
     f"{BUILDER_URL} . Without typing anything, click 'Download Word Agenda'.",
     "A document downloads. Opened in Word it says 'Untitled event' and "
     "'No sessions have been added yet.' It must not crash or produce a broken file."),
    ("A58", "Edge", "Still in the incognito window, paste this into a session Link box: "
     "javascript:alert(1) . Then add a date and title and download the file.",
     "The document downloads normally and that text does NOT appear as a clickable link in it. "
     "Nothing pops up. Report exactly what you see in the document."),

    # --- Responsive ---------------------------------------------------------
    ("A59", "Mobile", f"On a real phone, open {TOOL_URL}",
     "Landing page fits the screen. No sideways scrolling. Text is readable without zooming."),
    ("A60", "Mobile", "On the phone, tap the hamburger (three lines) menu in the top right.",
     "A menu opens with How it works, Features, Pricing, FAQ, Blog, News, Log in, Start for free. "
     "Tapping one closes the menu and goes there."),
    ("A61", "Mobile", "On the phone, build a small agenda: event name, one date, two sessions.",
     "All boxes are tappable, the keyboard does not cover what you are typing, and nothing "
     "overlaps or runs off the side of the screen."),
    ("A62", "Mobile", "On the phone, use the up and down chevrons to reorder the two sessions.",
     "They reorder correctly by tapping. (Dragging on a touch screen may be fiddly, that is what "
     "the arrows are for. Note in your answer whether dragging worked on your phone.)"),
    ("A63", "Mobile", "On the phone, download the agenda.",
     "The .docx downloads to the phone and can be opened."),
    ("A64", "Browsers", "Repeat rows A21, A29 and A44 in a second browser "
     "(Firefox or Safari, whichever you did not use).",
     "Adding a session, dragging to reorder, and downloading all behave the same as in Chrome."),
]

# Phase B is gated. The conversion handoff is not built yet, so there is
# deliberately nothing here that needs a production login.
ROWS_B = [
    ("B1", "SCOTT", "Scott: create a throwaway test account and put its email and password here "
     "before sending this to the tester.", "Not started until Scott says so."),
    ("B2", "Account", "Log in with the test account, then go to "
     f"{BUILDER_URL} and build a short agenda.",
     "The builder works the same when logged in, and still shows the marketing style header."),
    ("B3", "Account", "Download the agenda while logged in.",
     "Downloads exactly as it does when logged out."),
    ("B4", "Account", "Click 'Create Your Event' while logged in.",
     "You are taken to the events page with the create form open. NOTE: the agenda details are "
     "NOT expected to carry across yet, that part is still being built. Just confirm where it goes."),
    ("B5", "Account", "Go back to the builder after that.",
     "Your agenda is still saved and unchanged."),
]


def build_media() -> None:
    MEDIA.mkdir(parents=True, exist_ok=True)

    square = Image.new("RGB", (400, 400), (37, 99, 235))
    d = ImageDraw.Draw(square)
    d.ellipse((90, 90, 310, 310), fill=(255, 255, 255))
    square.save(MEDIA / "logo-square.png")

    wide = Image.new("RGB", (600, 200), (13, 27, 42))
    d = ImageDraw.Draw(wide)
    d.rectangle((40, 60, 560, 140), fill=(79, 141, 247))
    wide.save(MEDIA / "logo-wide.jpg", format="JPEG", quality=90)

    trans = Image.new("RGBA", (300, 300), (0, 0, 0, 0))
    d = ImageDraw.Draw(trans)
    d.ellipse((30, 30, 270, 270), fill=(37, 99, 235, 255))
    trans.save(MEDIA / "logo-transparent.webp", format="WEBP")

    # Comfortably over the 1 MB cap: random noise does not compress away.
    import os

    noise = Image.frombytes("RGB", (900, 900), os.urandom(900 * 900 * 3))
    noise.save(MEDIA / "oversized-logo.png", format="PNG", compress_level=0)

    (MEDIA / "not-an-image.png").write_text(
        "This is a plain text file that has been given a .png name on purpose.\n"
        "It is here so you can check the site refuses files that are not really images.\n",
        encoding="utf-8",
    )

    size_mb = (MEDIA / "oversized-logo.png").stat().st_size / (1024 * 1024)
    assert size_mb > 1, f"oversized-logo.png is only {size_mb:.2f} MB, it must exceed 1 MB"


def build_workbook() -> None:
    wb = Workbook()

    ws = wb.active
    ws.title = "READ ME"
    lines = [
        ("Intro Connect: Agenda Builder QA", True),
        ("", False),
        ("Fill in ONLY the three yellow columns on the Test Script sheet:", False),
        ("  Pass/Fail (dropdown: PASS / FAIL / BLOCKED)", False),
        ("  What actually happened (one or two sentences; exact error text if any)", False),
        ("  Screenshot file(s) (filename in your screenshots folder, e.g. A29.png)", False),
        ("Everything else is read only reference. Work top to bottom, keep the order.", False),
        ("", False),
        ("PASS means it did exactly what the 'What should happen' column says.", False),
        ("FAIL means it did something else. Say what it actually did.", False),
        ("BLOCKED means you could not attempt the row. Say why.", False),
        ("", False),
        ("Example of a filled row:", True),
        ('  A29 | FAIL | "Card followed the mouse but snapped back to its old place on release." | A29.png',
         False),
        ("", False),
        ("Screenshot every FAIL. Screenshots are optional on a PASS unless the row asks.", False),
        ("", False),
        ("Do Phase A only. Scott will tell you if Phase B is a go.", False),
        ("", False),
        ("Full instructions and the rules: QA-TESTER-HANDOUT.docx in this package.", False),
    ]
    for i, (text, bold) in enumerate(lines, start=1):
        cell = ws.cell(row=i, column=1, value=text)
        cell.font = Font(bold=bold, size=12 if bold and i == 1 else 11)
    ws.column_dimensions["A"].width = 105

    sh = wb.create_sheet("Test Script")
    head_fill = PatternFill("solid", fgColor="0D1B2A")
    fill_fill = PatternFill("solid", fgColor="FFF3C4")
    phase_fill = PatternFill("solid", fgColor="E8EEF9")

    for ci, name in enumerate(HEADER, start=1):
        c = sh.cell(row=1, column=ci, value=name)
        c.font = Font(bold=True, color="FFFFFF")
        c.fill = head_fill
        c.alignment = Alignment(vertical="center", wrap_text=True)

    r = 2

    def add_banner(text: str):
        nonlocal r
        c = sh.cell(row=r, column=1, value=text)
        c.font = Font(bold=True)
        c.fill = phase_fill
        sh.merge_cells(start_row=r, start_column=1, end_row=r, end_column=len(HEADER))
        r += 1

    def add_rows(rows):
        nonlocal r
        for row_id, area, todo, expect in rows:
            for ci, value in enumerate([row_id, area, todo, expect], start=1):
                c = sh.cell(row=r, column=ci, value=value)
                c.alignment = Alignment(vertical="top", wrap_text=True)
            for ci in range(5, len(HEADER) + 1):
                c = sh.cell(row=r, column=ci)
                c.fill = fill_fill
                c.alignment = Alignment(vertical="top", wrap_text=True)
            r += 1

    add_banner("PHASE A  —  the free tool, no account needed.  Do this whole phase.")
    add_rows(ROWS)
    add_banner("PHASE B  —  logged in checks.  DO NOT START until Scott confirms.")
    add_rows(ROWS_B)

    widths = [7, 13, 62, 62, 11, 42, 20]
    for ci, w in enumerate(widths, start=1):
        sh.column_dimensions[get_column_letter(ci)].width = w
    sh.freeze_panes = "A2"

    dv = DataValidation(type="list", formula1='"PASS,FAIL,BLOCKED"', allow_blank=True)
    sh.add_data_validation(dv)
    dv.add(f"E2:E{r - 1}")

    sh.auto_filter.ref = f"A1:{get_column_letter(len(HEADER))}{r - 1}"

    wb.save(OUT / "QA-TEST-SCRIPT.xlsx")


def _p(doc, text="", *, bold=False, size=11, color=None, space_after=6, align=None):
    para = doc.add_paragraph()
    if align:
        para.alignment = align
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return para


def build_handout() -> None:
    doc = Document()
    doc.core_properties.title = "Intro Connect Agenda Builder QA"

    _p(doc, "Intro Connect", bold=True, size=20, color=RGBColor(0x1D, 0x4E, 0xD8), space_after=0)
    _p(doc, "Agenda Builder: manual QA", bold=True, size=15, space_after=14)

    _p(doc, "What this is", bold=True, size=13)
    _p(doc, "Intro Connect helps event organizers turn an event into a private directory of "
            "everyone who attended, so people can actually find each other afterwards.")
    _p(doc, "The Agenda Builder is a new free tool. Anyone can build an event agenda in the "
            "browser and download it as a Word document, with no account and no payment. It is "
            "brand new and has never been tested by anyone but the person who built it. That is "
            "your job.", space_after=14)

    _p(doc, "Where to go", bold=True, size=13)
    _p(doc, TOOL_URL, bold=True)
    _p(doc, "Nothing on the public website links to this page yet, so the address above is the "
            "only way in. That is on purpose.", space_after=14)

    _p(doc, "You do not need a login", bold=True, size=13)
    _p(doc, "The whole of Phase A works without an account, and you should not create one. If any "
            "row seems to ask you to sign up, re-read it: the only row that touches sign up is A55, "
            "and it asks you to look at where the button goes and then press Back.", space_after=14)

    _p(doc, "The job", bold=True, size=13)
    _p(doc, "Open QA-TEST-SCRIPT.xlsx and work through the Test Script tab top to bottom, in "
            "order. Fill in the three yellow columns as you go. Do Phase A only. Phase B is greyed "
            "out until Scott says otherwise.")
    _p(doc, "Estimated time: 4 to 6 hours.", bold=True, space_after=14)

    _p(doc, "The rules that matter", bold=True, size=13)
    rules = [
        "This is a live production site. Do not create accounts, do not sign up, and do not "
        "enter any real personal details anywhere. The tool needs none of it.",
        "Use made up event details. 'Denver Founders Dinner' and similar are fine. Do not use "
        "real people's names, real email addresses, or anything you would not want in a test log.",
        "All uploads come from the sample-media folder in this package. Do not upload your own "
        "photos or any real company logo.",
        "Work in order and do not skip rows. Several rows depend on what the row before it set up.",
        "If something is unclear, or you hit an error you cannot get past, or you see anything "
        "that looks like a real person's data, stop and message Scott rather than improvising.",
        "Everything in this package, and everything you see on the site, is confidential. Do not "
        "share it, post it, or feed it into anything.",
    ]
    for rule in rules:
        para = doc.add_paragraph(style="List Number")
        para.paragraph_format.space_after = Pt(6)
        para.add_run(rule).font.size = Pt(11)
    doc.add_paragraph()

    _p(doc, "Screenshots", bold=True, size=13)
    _p(doc, "Make a folder called screenshots next to the workbook. Name each file after its row: "
            "A29.png, and A29-2.png if you need a second. Screenshot every FAIL. On a PASS a "
            "screenshot is optional unless the row asks for one.", space_after=14)

    _p(doc, "What good feedback looks like", bold=True, size=13)
    _p(doc, "Not helpful:  \"Doesn't work\"", space_after=2)
    _p(doc, "Helpful:  \"Dragged the first session below the last one. The card followed the mouse "
            "but snapped back to its original position when I released. Happened every time in "
            "Chrome, worked in Firefox.\"", space_after=14)

    _p(doc, "A few things worth knowing", bold=True, size=13)
    notes = [
        "Times are 24 hour in the builder (17:30) and become 12 hour in the document (5:30 PM). "
        "That is intended.",
        "Sessions stay in the order YOU put them in. The tool should never silently re-sort them "
        "by time. There is a separate 'Sort by time' button for that. Several rows check this.",
        "Overlapping sessions give an orange warning but must never block you. A session that ends "
        "before it starts is a red error and does block the download. Both are intended.",
        "'Private notes' on a session must never appear in the downloaded document. Row A51 checks "
        "this and it is one of the most important rows in the script.",
        "Your work is saved in the browser itself, not on a server, so it survives a refresh but "
        "will not follow you to a different computer or an incognito window.",
    ]
    for note in notes:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(6)
        para.add_run(note).font.size = Pt(11)
    doc.add_paragraph()

    _p(doc, "When you are done", bold=True, size=13)
    _p(doc, "Zip the filled in workbook together with your screenshots folder and send it back to "
            "Scott the same way you received this package.")

    doc.save(OUT / "QA-TESTER-HANDOUT.docx")


README = f"""INTRO CONNECT: AGENDA BUILDER QA PACKAGE — START HERE

1. Open QA-TESTER-HANDOUT.docx and read it first. It explains what you are
   testing and the rules that matter.

2. Open QA-TEST-SCRIPT.xlsx. Start at the READ ME tab, then work through the
   Test Script tab top to bottom, in order. Do Phase A only; Scott will tell
   you when Phase B is a go.

3. sample-media\\ holds every file the test rows ask you to upload. Use only
   these files for uploads.

4. Create a folder called "screenshots" next to the workbook. Save every
   screenshot there, named after the row (A29.png, A29-2.png, ...).

5. When Phase A is done: zip the filled workbook + your screenshots folder and
   send it back to Scott the same way you received this package.

The tool under test is here, and nothing links to it yet:
   {TOOL_URL}

You do NOT need an account and you should not create one.

Questions or anything unexpected (an error you cannot get past, anything that
looks like real user data): stop and contact Scott before continuing.

This package is confidential. Do not share any part of it.
"""


def main() -> None:
    if OUT.exists():
        shutil.rmtree(OUT)
    OUT.mkdir(parents=True)

    build_media()
    build_workbook()
    build_handout()
    (OUT / "README-package.txt").write_text(README, encoding="utf-8")

    zip_path = ROOT / "build" / "QA-AGENDA-BUILDER-PACKAGE.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
        for path in sorted(OUT.rglob("*")):
            if path.is_file():
                z.write(path, path.relative_to(OUT))

    total_rows = len(ROWS) + len(ROWS_B)
    print(f"wrote {zip_path} ({zip_path.stat().st_size / 1024:.0f} KB)")
    print(f"  test rows: {len(ROWS)} phase A + {len(ROWS_B)} phase B = {total_rows}")
    for path in sorted(OUT.rglob("*")):
        if path.is_file():
            print(f"  {path.relative_to(OUT)}  ({path.stat().st_size / 1024:.1f} KB)")


if __name__ == "__main__":
    main()
