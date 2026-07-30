"""Agenda export endpoint, exercised through the real FastAPI app.

Like tests/test_route_auth.py, the TestClient is used WITHOUT its context
manager so the app lifespan (Mongo index creation) never runs. The export route
is stateless and never touches the database, so no live MongoDB is needed.

Run from backend/: python -m pytest tests/test_agenda_api.py
"""
import io
import os
from datetime import date

os.environ.setdefault("MONGO_URL", "mongodb://localhost:27017")
os.environ.setdefault("DB_NAME", "t")
os.environ.setdefault("JWT_SECRET", "x")

import pytest
from docx import Document
from fastapi.testclient import TestClient
from PIL import Image

import rate_limit
import server
from routers.agenda import DOCX_MEDIA_TYPE, EXPORT_LIMIT

client = TestClient(server.app)

EXPORT = "/api/agenda/export"


@pytest.fixture(autouse=True)
def _clear_rate_limit():
    """The limiter is process global, so a test that exhausts the budget would
    otherwise 429 every later test in the session."""
    rate_limit._hits.clear()
    yield
    rate_limit._hits.clear()


def _agenda(**over):
    base = {
        "event_name": "Denver Founders Dinner",
        "start_date": "2026-08-01",
        "items": [
            {
                "date": "2026-08-01",
                "start_time": "18:00",
                "end_time": "19:00",
                "title": "Welcome and introductions",
            }
        ],
    }
    base.update(over)
    return base


def _png_data_url() -> str:
    import base64

    buf = io.BytesIO()
    Image.new("RGB", (40, 40), (37, 99, 235)).save(buf, format="PNG")
    return "data:image/png;base64," + base64.b64encode(buf.getvalue()).decode()


def test_export_is_reachable_without_a_session():
    """The builder must work before the visitor has an account."""
    res = client.post(EXPORT, json=_agenda())
    assert res.status_code == 200


def test_export_returns_a_docx_attachment():
    res = client.post(EXPORT, json=_agenda())
    assert res.headers["content-type"] == DOCX_MEDIA_TYPE
    assert res.headers["content-disposition"] == (
        'attachment; filename="denver-founders-dinner-agenda.docx"'
    )
    assert res.content[:2] == b"PK"


def test_exported_document_contains_the_session():
    res = client.post(EXPORT, json=_agenda())
    doc = Document(io.BytesIO(res.content))
    text = "\n".join(
        p.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
        for p in cell.paragraphs
    )
    assert "Welcome and introductions" in text
    assert "6:00 PM to 7:00 PM" in text


def test_export_accepts_a_logo():
    res = client.post(EXPORT, json=_agenda(logo=_png_data_url()))
    assert res.status_code == 200
    # A document carrying an image is meaningfully larger than one without.
    plain = client.post(EXPORT, json=_agenda())
    assert len(res.content) > len(plain.content)


def test_export_rejects_a_logo_that_is_not_an_image():
    import base64

    bad = "data:image/png;base64," + base64.b64encode(b"not an image").decode()
    res = client.post(EXPORT, json=_agenda(logo=bad))
    assert res.status_code == 400


def test_export_rejects_reversed_session_times():
    bad = _agenda(
        items=[{
            "date": "2026-08-01",
            "start_time": "19:00",
            "end_time": "18:00",
            "title": "Backwards",
        }]
    )
    assert client.post(EXPORT, json=bad).status_code == 422


def _document_text(blob: bytes) -> str:
    doc = Document(io.BytesIO(blob))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


def test_export_strips_a_dangerous_link_but_still_succeeds():
    """A bad link is dropped, not fatal. Asserted against the extracted
    document text: a .docx is a zip, so scanning the raw response bytes would
    pass even when the string is present."""
    payload = _agenda(
        items=[{
            "date": "2026-08-01",
            "start_time": "18:00",
            "end_time": "19:00",
            "title": "Talk",
            "external_url": "javascript:alert(1)",
        }]
    )
    res = client.post(EXPORT, json=payload)
    assert res.status_code == 200
    text = _document_text(res.content)
    assert "javascript" not in text.lower()
    assert "Talk" in text


def test_export_keeps_a_safe_link_in_the_document():
    """Counterpart to the test above: proves the assertion can actually fail,
    rather than passing because links are never rendered at all."""
    payload = _agenda(
        items=[{
            "date": "2026-08-01",
            "start_time": "18:00",
            "end_time": "19:00",
            "title": "Talk",
            "external_url": "https://example.com/talk",
        }]
    )
    text = _document_text(client.post(EXPORT, json=payload).content)
    assert "https://example.com/talk" in text


def test_landing_page_is_public_and_serves_html():
    res = client.get("/agenda")
    assert res.status_code == 200
    assert "text/html" in res.headers["content-type"]
    assert "Free event agenda builder" in res.text


def test_landing_page_is_cacheable():
    """It is static content behind a proxy; without this every hit pays the
    Render cold start."""
    res = client.get("/agenda")
    assert "s-maxage" in res.headers.get("cache-control", "")


def test_export_is_rate_limited_per_ip():
    payload = _agenda()
    for _ in range(EXPORT_LIMIT):
        assert client.post(EXPORT, json=payload).status_code == 200
    assert client.post(EXPORT, json=payload).status_code == 429
