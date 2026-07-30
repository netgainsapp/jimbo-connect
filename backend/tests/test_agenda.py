"""Agenda Builder: schema validation and Word rendering.

Pure unit tests, no DB and no network. Run from backend/:
    python -m pytest tests/test_agenda.py
"""
import io
from datetime import date

import pytest
from docx import Document
from pydantic import ValidationError

from agenda.docx import (
    build_docx,
    date_range_line,
    format_date,
    format_time,
    format_time_range,
)
from agenda.schema import (
    MAX_ITEMS,
    AgendaExportRequest,
    AgendaItem,
    clean_text,
    clean_url,
    group_by_day,
    slugify_filename,
)


def _item(**over):
    base = {
        "date": date(2026, 8, 1),
        "start_time": "09:00",
        "end_time": "10:00",
        "title": "Opening remarks",
    }
    base.update(over)
    return AgendaItem(**base)


def _docx_text(blob: bytes) -> str:
    """All visible text in the document, paragraphs and table cells."""
    doc = Document(io.BytesIO(blob))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


# --------------------------------------------------------------------------
# URL handling: these become text in a document that travels by email.
# --------------------------------------------------------------------------

def test_drops_javascript_scheme_url():
    assert clean_url("javascript:alert(1)") == ""


def test_drops_data_and_file_scheme_urls():
    assert clean_url("data:text/html;base64,PHNjcmlwdD4=") == ""
    assert clean_url("file:///etc/passwd") == ""


def test_drops_scheme_smuggled_through_interior_whitespace():
    """Interior whitespace is discarded by consumers, so a validator that only
    trims the ends can pass a string that reconstitutes into javascript:."""
    assert clean_url("java\nscript:alert(1)") == ""
    assert clean_url("java\tscript:alert(1)") == ""
    assert clean_url("  javascript :alert(1)") == ""


def test_drops_uppercase_scheme_variants():
    assert clean_url("JavaScript:alert(1)") == ""
    assert clean_url("DATA:text/html,x") == ""


def test_keeps_https_url():
    assert clean_url("https://example.com/talk") == "https://example.com/talk"


def test_bare_domain_is_promoted_to_https():
    assert clean_url("example.com/talk") == "https://example.com/talk"


def test_bad_item_url_is_dropped_without_failing_the_whole_agenda():
    """A single bad link must not cost the organizer their export."""
    item = _item(external_url="javascript:alert(1)")
    assert item.external_url == ""
    assert item.title == "Opening remarks"


# --------------------------------------------------------------------------
# Text and time validation
# --------------------------------------------------------------------------

def test_strips_control_characters_from_text():
    assert clean_text("Ope\x00ning\x07 remarks") == "Opening remarks"


def test_keeps_newlines_in_descriptions():
    assert clean_text("line one\nline two") == "line one\nline two"


def test_rejects_end_time_before_start_time():
    with pytest.raises(ValidationError):
        _item(start_time="14:00", end_time="09:00")


def test_rejects_end_time_equal_to_start_time():
    with pytest.raises(ValidationError):
        _item(start_time="09:00", end_time="09:00")


def test_rejects_malformed_time():
    with pytest.raises(ValidationError):
        _item(start_time="9am")
    with pytest.raises(ValidationError):
        _item(start_time="25:00")


def test_allows_item_with_no_times():
    item = _item(start_time="", end_time="")
    assert item.start_time == ""


def test_rejects_end_date_before_start_date():
    with pytest.raises(ValidationError):
        AgendaExportRequest(
            event_name="Summit",
            start_date=date(2026, 8, 5),
            end_date=date(2026, 8, 1),
        )


def test_rejects_more_items_than_the_cap():
    too_many = [_item() for _ in range(MAX_ITEMS + 1)]
    with pytest.raises(ValidationError):
        AgendaExportRequest(event_name="Summit", items=too_many)


# --------------------------------------------------------------------------
# Grouping and ordering
# --------------------------------------------------------------------------

def test_groups_days_chronologically():
    items = [
        _item(date=date(2026, 8, 2), start_time="09:00", end_time="", title="Day two"),
        _item(date=date(2026, 8, 1), start_time="14:00", end_time="", title="Day one"),
    ]
    groups = group_by_day(items)
    assert [d for d, _ in groups] == [date(2026, 8, 1), date(2026, 8, 2)]


def test_preserves_manual_order_within_a_day_even_when_times_disagree():
    """Regression: an earlier version sorted each day by start time, which made
    the builder's reorder controls a no-op and let the document disagree with
    the preview. The organizer's arrangement wins."""
    items = [
        _item(date=date(2026, 8, 1), start_time="17:00", end_time="", title="Closing"),
        _item(date=date(2026, 8, 1), start_time="09:00", end_time="", title="Opening"),
    ]
    _, ordered = group_by_day(items)[0]
    assert [i.title for i in ordered] == ["Closing", "Opening"]


def test_preserves_position_of_untimed_sessions():
    items = [
        _item(start_time="", end_time="", title="Unscheduled"),
        _item(start_time="09:00", end_time="", title="Scheduled"),
    ]
    _, ordered = group_by_day(items)[0]
    assert [i.title for i in ordered] == ["Unscheduled", "Scheduled"]


def test_exported_document_follows_manual_order():
    """End to end guard: the docx must render the arranged order, not a
    time sort."""
    agenda = AgendaExportRequest(
        event_name="Summit",
        items=[
            _item(date=date(2026, 8, 1), start_time="17:00", end_time="", title="Closing keynote"),
            _item(date=date(2026, 8, 1), start_time="09:00", end_time="", title="Opening keynote"),
        ],
    )
    text = _docx_text(build_docx(agenda))
    assert text.index("Closing keynote") < text.index("Opening keynote")


# --------------------------------------------------------------------------
# Formatting helpers
# --------------------------------------------------------------------------

def test_formats_date_with_weekday_and_no_zero_padding():
    assert format_date(date(2026, 8, 1)) == "Saturday, August 1, 2026"


def test_formats_twelve_hour_times():
    assert format_time("09:00") == "9:00 AM"
    assert format_time("13:30") == "1:30 PM"
    assert format_time("00:15") == "12:15 AM"
    assert format_time("12:00") == "12:00 PM"
    assert format_time("") == ""


def test_time_range_falls_back_to_whichever_end_exists():
    assert format_time_range("09:00", "10:00") == "9:00 AM to 10:00 AM"
    assert format_time_range("09:00", "") == "9:00 AM"
    assert format_time_range("", "") == ""


def test_date_range_falls_back_to_the_span_of_the_sessions():
    agenda = AgendaExportRequest(
        event_name="Summit",
        items=[_item(date=date(2026, 8, 1)), _item(date=date(2026, 8, 3))],
    )
    assert date_range_line(agenda) == (
        "Saturday, August 1, 2026 to Monday, August 3, 2026"
    )


def test_location_line_prefers_venue_and_falls_back_to_virtual_url():
    with_venue = AgendaExportRequest(
        event_name="S", venue_name="The Loft", venue_address="12 Main St",
        virtual_url="https://zoom.example.com/x",
    )
    assert with_venue.location_line() == "The Loft, 12 Main St"

    virtual = AgendaExportRequest(
        event_name="S", virtual_url="https://zoom.example.com/x"
    )
    assert virtual.location_line() == "https://zoom.example.com/x"


def test_slugify_filename_is_ascii_and_bounded():
    assert slugify_filename("Denver Founders Dinner!") == "denver-founders-dinner"
    assert slugify_filename("") == "agenda"
    assert len(slugify_filename("x" * 200)) <= 60


# --------------------------------------------------------------------------
# Word rendering
# --------------------------------------------------------------------------

def test_export_produces_a_real_docx_zip():
    blob = build_docx(AgendaExportRequest(event_name="Summit"))
    assert blob[:2] == b"PK", "a .docx is an Open XML zip"
    assert len(blob) > 1000


def test_document_contains_event_details_and_sessions():
    agenda = AgendaExportRequest(
        event_name="Denver Founders Dinner",
        description="An evening for operators.",
        venue_name="The Loft",
        organizer_name="Scott Weiss",
        organizer_company="Intro Connect",
        items=[_item(title="Opening remarks", speaker="Scott", location="Main room")],
    )
    text = _docx_text(build_docx(agenda))
    for expected in (
        "Denver Founders Dinner",
        "An evening for operators.",
        "The Loft",
        "Saturday, August 1, 2026",
        "9:00 AM to 10:00 AM",
        "Opening remarks",
        "Scott",
        "Main room",
    ):
        assert expected in text, f"missing from document: {expected}"


def test_private_notes_are_never_rendered_into_the_export():
    agenda = AgendaExportRequest(
        event_name="Summit",
        items=[_item(notes="remember to pay the caterer")],
    )
    assert "caterer" not in _docx_text(build_docx(agenda))


def test_document_default_font_is_calibri():
    """Asserted on the rFonts slots, not just style.font.name: Word resolves
    typefaces through rFonts, so a document can carry the right style name and
    still render in a substituted font."""
    from docx.oxml.ns import qn

    from agenda.docx import BODY_FONT

    doc = Document(io.BytesIO(build_docx(AgendaExportRequest(event_name="Summit"))))
    normal = doc.styles["Normal"]
    assert normal.font.name == BODY_FONT

    rfonts = normal.element.get_or_add_rPr().get_or_add_rFonts()
    for slot in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        assert rfonts.get(qn(slot)) == BODY_FONT, f"{slot} not set to {BODY_FONT}"


def test_document_carries_the_footer_credit():
    doc = Document(io.BytesIO(build_docx(AgendaExportRequest(event_name="Summit"))))
    footer_text = "\n".join(p.text for p in doc.sections[0].footer.paragraphs)
    assert "Intro Connect" in footer_text
    assert "intro-connect.com" in footer_text


def test_empty_agenda_still_exports():
    text = _docx_text(build_docx(AgendaExportRequest(event_name="")))
    assert "Untitled event" in text
    assert "No sessions have been added yet." in text


# --------------------------------------------------------------------------
# Landing page (the SEO surface)
# --------------------------------------------------------------------------

def test_landing_renders_real_content_without_javascript():
    """The whole point of server rendering this page: a crawler that runs no JS
    must still see the copy."""
    from agenda import landing

    out = landing.render_landing()
    assert "<h1" in out
    assert "Free event agenda builder" in out
    assert "How it works" in out
    assert "Common questions" in out
    assert len(out) > 4000


def test_landing_points_at_the_builder():
    from agenda import landing

    out = landing.render_landing()
    assert landing.BUILDER_URL in out
    assert out.count(landing.BUILDER_URL) >= 2, "expected a CTA near the top and the bottom"


def test_landing_declares_a_free_web_application():
    from agenda import landing

    out = landing.render_landing()
    # seo.json_ld serialises compactly (separators=(",", ":")), so there is no
    # space after the colon.
    assert '"@type":"WebApplication"' in out
    assert '"price":"0"' in out


def test_landing_faq_markup_matches_the_visible_copy():
    """Search engines only honour FAQPage markup when the same text is on the
    page, so assert both rather than just the JSON-LD."""
    from agenda import landing

    out = landing.render_landing()
    assert '"@type":"FAQPage"' in out
    for question, answer in landing.FAQS:
        assert question in out
        assert answer[:40] in out


def test_landing_sets_its_own_canonical():
    from agenda import landing

    out = landing.render_landing()
    assert 'rel="canonical"' in out
    assert "/agenda" in out


def test_multi_day_agenda_renders_a_heading_per_day():
    agenda = AgendaExportRequest(
        event_name="Two Day Summit",
        items=[
            _item(date=date(2026, 8, 1), title="Day one talk"),
            _item(date=date(2026, 8, 2), title="Day two talk"),
        ],
    )
    text = _docx_text(build_docx(agenda))
    assert "Saturday, August 1, 2026" in text
    assert "Sunday, August 2, 2026" in text
