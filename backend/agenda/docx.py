"""Word (.docx) rendering for an agenda.

Produces a real Open XML document rather than HTML served with a .doc
extension, so it stays genuinely editable after download. Layout is
deliberately restrained: a small logo, clear day headings, and a two column
time/session table that survives being edited and reflows sensibly when
printed. Per the spec this carries a modest footer credit and no advertising.
"""
from __future__ import annotations

import io
from datetime import date as _date
from typing import Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from .schema import AgendaExportRequest, AgendaItem, group_by_day

# Restrained ink. Close to the product's primary (#2563EB) without turning the
# document into a brand asset.
HEADING_COLOR = RGBColor(0x1D, 0x4E, 0xD8)
MUTED_COLOR = RGBColor(0x6B, 0x72, 0x80)
BODY_COLOR = RGBColor(0x0A, 0x0C, 0x10)

LOGO_MAX_WIDTH_INCHES = 1.4
FOOTER_TEXT = "Agenda created with Intro Connect"
FOOTER_URL = "intro-connect.com"

_MONTHS = (
    "January", "February", "March", "April", "May", "June",
    "July", "August", "September", "October", "November", "December",
)
_DAYS = (
    "Monday", "Tuesday", "Wednesday", "Thursday",
    "Friday", "Saturday", "Sunday",
)


def format_date(value: _date, *, with_weekday: bool = True) -> str:
    """"Saturday, August 1, 2026". Built by hand because platform strftime
    differs on zero padding (%-d is not portable to Windows)."""
    if not value:
        return ""
    stem = f"{_MONTHS[value.month - 1]} {value.day}, {value.year}"
    return f"{_DAYS[value.weekday()]}, {stem}" if with_weekday else stem


def format_time(value: str) -> str:
    """"09:00" -> "9:00 AM". Empty input stays empty."""
    if not value:
        return ""
    hour, minute = int(value[:2]), value[3:5]
    suffix = "AM" if hour < 12 else "PM"
    display = hour % 12 or 12
    return f"{display}:{minute} {suffix}"


def format_time_range(start: str, end: str) -> str:
    if start and end:
        return f"{format_time(start)} to {format_time(end)}"
    return format_time(start) or format_time(end)


def date_range_line(agenda: AgendaExportRequest) -> str:
    start, end = agenda.start_date, agenda.end_date
    if start and end and end != start:
        return f"{format_date(start)} to {format_date(end)}"
    if start:
        return format_date(start)
    # Fall back to the span the sessions themselves cover.
    days = sorted({i.date for i in agenda.items})
    if not days:
        return ""
    if len(days) == 1:
        return format_date(days[0])
    return f"{format_date(days[0])} to {format_date(days[-1])}"


def _run(paragraph, text, *, bold=False, size=11, color=BODY_COLOR, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return run


def _add_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(para, f"{FOOTER_TEXT}   |   {FOOTER_URL}", size=8, color=MUTED_COLOR)


def _add_logo(doc: Document, logo_png: Optional[bytes]) -> None:
    if not logo_png:
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(io.BytesIO(logo_png), width=Inches(LOGO_MAX_WIDTH_INCHES))


def _add_session_row(table, item: AgendaItem) -> None:
    row = table.add_row()
    time_cell, detail_cell = row.cells[0], row.cells[1]

    time_para = time_cell.paragraphs[0]
    _run(time_para, format_time_range(item.start_time, item.end_time) or "TBD",
         bold=True, size=10)

    detail_para = detail_cell.paragraphs[0]
    _run(detail_para, item.title or "Untitled session", bold=True, size=11)

    meta = " | ".join(p for p in (
        f"Speaker: {item.speaker}" if item.speaker else "",
        item.location,
    ) if p)
    if meta:
        para = detail_cell.add_paragraph()
        _run(para, meta, size=9, color=MUTED_COLOR)

    if item.description:
        para = detail_cell.add_paragraph()
        _run(para, item.description, size=10)

    if item.external_url:
        para = detail_cell.add_paragraph()
        # Written as plain text rather than a Word hyperlink field. The scheme
        # is already restricted to http(s) at the schema boundary; keeping it
        # inert here means a document that travels by email carries no
        # clickable target we did not construct ourselves.
        _run(para, item.external_url, size=9, color=MUTED_COLOR)


def _add_day(doc: Document, day: _date, items: list) -> None:
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(16)
    heading.paragraph_format.space_after = Pt(6)
    _run(heading, format_date(day), bold=True, size=13, color=HEADING_COLOR)

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for item in items:
        _add_session_row(table, item)
    # Narrow time column, wide detail column. Set per cell: Word honours cell
    # widths more reliably than column widths when the file is later edited.
    for row in table.rows:
        row.cells[0].width = Inches(1.35)
        row.cells[1].width = Inches(5.15)


def _add_organizer(doc: Document, agenda: AgendaExportRequest) -> None:
    lines = [p for p in (
        agenda.organizer_name,
        agenda.organizer_company,
        agenda.organizer_email,
        agenda.event_website,
    ) if p]
    if not lines:
        return
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(18)
    _run(heading, "Organizer", bold=True, size=11, color=HEADING_COLOR)
    for line in lines:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        _run(para, line, size=10, color=MUTED_COLOR)


def build_docx(agenda: AgendaExportRequest, logo_png: Optional[bytes] = None) -> bytes:
    """Render the agenda and return the .docx bytes."""
    doc = Document()
    doc.core_properties.title = agenda.display_name()
    if agenda.organizer_name:
        doc.core_properties.author = agenda.organizer_name

    _add_logo(doc, logo_png)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    _run(title, agenda.display_name(), bold=True, size=22)

    for line in (date_range_line(agenda), agenda.location_line()):
        if not line:
            continue
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(0)
        _run(para, line, size=10, color=MUTED_COLOR)

    if agenda.description:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(14)
        _run(para, agenda.description, size=11)

    groups = group_by_day(agenda.items)
    if groups:
        for day, items in groups:
            _add_day(doc, day, items)
    else:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(14)
        _run(para, "No sessions have been added yet.", size=10,
             color=MUTED_COLOR, italic=True)

    _add_organizer(doc, agenda)
    _add_footer(doc)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
